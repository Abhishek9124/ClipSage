import os
import re
import json
import hashlib
from pathlib import Path

from dotenv import load_dotenv
import streamlit as st

from youtube_transcript_api import YouTubeTranscriptApi

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
from langchain_chroma import Chroma
from langchain_core.prompts import ChatPromptTemplate

load_dotenv()

CACHE_DIR = Path(".clipsage_cache")
CACHE_DIR.mkdir(exist_ok=True)
VECTOR_DIR = CACHE_DIR / "vectors"
VECTOR_DIR.mkdir(exist_ok=True)
TRANSCRIPT_DIR = CACHE_DIR / "transcripts"
TRANSCRIPT_DIR.mkdir(exist_ok=True)

llm = ChatGoogleGenerativeAI(model="gemini-2.5-flash", temperature=0.2)


def _run(prompt_template: str, **kwargs) -> str:
    prompt = ChatPromptTemplate.from_template(prompt_template)
    chain = prompt | llm
    return chain.invoke(kwargs).content


def extract_video_id(url: str):
    match = re.search(r"(?:v=|\/)([0-9A-Za-z_-]{11}).*", url)
    if match:
        return match.group(1)
    st.error("Invalid YouTube URL. Please enter a valid video link.")
    return None


def _transcript_cache_path(video_id: str) -> Path:
    return TRANSCRIPT_DIR / f"{video_id}.json"


def fetch_transcript_segments(video_id: str, language: str = "auto"):
    """Returns (segments, language_used). Segments: list of {text, start, duration}."""
    cache_file = _transcript_cache_path(video_id)
    if cache_file.exists():
        data = json.loads(cache_file.read_text(encoding="utf-8"))
        return data["segments"], data["language"]

    api = YouTubeTranscriptApi()
    try:
        if language == "auto":
            transcript_list = api.list(video_id)
            t = next(iter(transcript_list))
            fetched = t.fetch()
            language_used = t.language_code
        else:
            fetched = api.fetch(video_id, languages=[language])
            language_used = language

        segments = [
            {"text": s.text, "start": float(s.start), "duration": float(s.duration)}
            for s in fetched
        ]
        cache_file.write_text(
            json.dumps({"segments": segments, "language": language_used}),
            encoding="utf-8",
        )
        return segments, language_used
    except Exception as e:
        st.error(f"Error fetching transcript: {e}")
        return None, None


def segments_to_text(segments) -> str:
    return " ".join(s["text"] for s in segments)


def segments_to_timestamped_text(segments) -> str:
    """Format segments with [HH:MM:SS] markers for the LLM to reference."""
    lines = []
    for s in segments:
        ts = _format_timestamp(s["start"])
        lines.append(f"[{ts}] {s['text']}")
    return "\n".join(lines)


def _format_timestamp(seconds: float) -> str:
    seconds = int(seconds)
    h, rem = divmod(seconds, 3600)
    m, s = divmod(rem, 60)
    if h:
        return f"{h:02d}:{m:02d}:{s:02d}"
    return f"{m:02d}:{s:02d}"


def timestamp_to_seconds(ts: str) -> int:
    parts = [int(p) for p in ts.split(":")]
    if len(parts) == 3:
        return parts[0] * 3600 + parts[1] * 60 + parts[2]
    if len(parts) == 2:
        return parts[0] * 60 + parts[1]
    return int(parts[0])


def linkify_timestamps(text: str, video_id: str) -> str:
    """Convert [HH:MM:SS] markers in markdown to clickable YouTube links."""
    def repl(m):
        ts = m.group(1)
        seconds = timestamp_to_seconds(ts)
        return f"[[{ts}]](https://www.youtube.com/watch?v={video_id}&t={seconds}s)"
    return re.sub(r"\[(\d{1,2}:\d{2}(?::\d{2})?)\]", repl, text)


def translate_transcript(transcript: str, target_language: str = "English") -> str:
    return _run(
        """You are an expert translator. Translate the transcript into {target_language},
        preserving meaning, tone, nuance, idioms, and the speaker's voice.
        Do not summarize.

        Transcript:
        {transcript}""",
        transcript=transcript,
        target_language=target_language,
    )


def get_important_topics(transcript: str) -> str:
    return _run(
        """Extract the 5 most important topics discussed.
        Output a numbered list. Concise. Only topics actually present.

        Transcript:
        {transcript}""",
        transcript=transcript,
    )


LENGTH_GUIDE = {
    "Short": "Roughly 8-12 bullet points covering only the essentials.",
    "Medium": "Roughly 15-25 bullets organized under 3-5 subheadings.",
    "Detailed": "Comprehensive notes with subheadings, examples, and key takeaways.",
}


def generate_notes(timestamped_transcript: str, length: str = "Medium") -> str:
    guide = LENGTH_GUIDE.get(length, LENGTH_GUIDE["Medium"])
    return _run(
        """You are an AI note-taker. Read the transcript (lines prefixed with [HH:MM:SS] timestamps)
        and produce structured notes.

        Length: {guide}

        Requirements:
        - Bulleted points grouped under markdown subheadings (##).
        - For each major bullet, include the most relevant timestamp marker in [HH:MM:SS] form
          taken directly from the transcript so the reader can jump to that moment.
        - Short, clear sentences. No invented content.

        Transcript:
        {transcript}""",
        transcript=timestamped_transcript,
        guide=guide,
    )


def generate_quiz(transcript: str, num_questions: int = 5) -> str:
    return _run(
        """Create a {n}-question quiz based ONLY on the transcript.
        Mix of multiple-choice (4 options) and short-answer.
        Format:
        **Q1.** question
        - A) ...
        - B) ...
        - C) ...
        - D) ...
        **Answer:** letter — short explanation

        Transcript:
        {transcript}""",
        transcript=transcript,
        n=num_questions,
    )


def generate_flashcards(transcript: str, num_cards: int = 10) -> str:
    return _run(
        """Create {n} flashcards from the transcript.
        Format each as:
        **Q:** question
        **A:** concise answer

        Transcript:
        {transcript}""",
        transcript=transcript,
        n=num_cards,
    )


def extract_key_quotes(timestamped_transcript: str, num_quotes: int = 7) -> str:
    return _run(
        """Extract the {n} most memorable or insightful quotes from the transcript.
        Each must be verbatim from the transcript.
        Format:
        - "quote text" — [HH:MM:SS]

        Transcript:
        {transcript}""",
        transcript=timestamped_transcript,
        n=num_quotes,
    )


def analyze_sentiment(transcript: str) -> str:
    return _run(
        """Analyze the sentiment and tone of this transcript.
        Provide:
        - **Overall sentiment:** (positive / negative / neutral / mixed) with confidence
        - **Dominant tone:** (e.g. educational, persuasive, humorous, critical)
        - **Emotional arc:** how the tone shifts across the video
        - **Notable emotional moments:** 2-3 examples

        Transcript:
        {transcript}""",
        transcript=transcript,
    )


def generate_blog_post(transcript: str) -> str:
    return _run(
        """Convert this transcript into a polished blog post (~800 words).
        Include: catchy title (# heading), intro hook, 3-5 sections with ## subheadings,
        and a conclusion. Use markdown.

        Transcript:
        {transcript}""",
        transcript=transcript,
    )


def generate_linkedin_post(transcript: str) -> str:
    return _run(
        """Write a LinkedIn post (~200 words) summarizing the key insight from this video.
        Hook in the first line. Short paragraphs. End with a question to drive engagement.
        Add 3-5 relevant hashtags at the end.

        Transcript:
        {transcript}""",
        transcript=transcript,
    )


def generate_twitter_thread(transcript: str) -> str:
    return _run(
        """Write a Twitter/X thread of 6-10 tweets summarizing this video.
        Each tweet under 280 characters. Number them (1/, 2/, ...).
        Tweet 1 must hook the reader. Last tweet has a takeaway + 2-3 hashtags.

        Transcript:
        {transcript}""",
        transcript=transcript,
    )


def create_chunks(transcript: str):
    splitter = RecursiveCharacterTextSplitter(chunk_size=10000, chunk_overlap=1000)
    return splitter.create_documents([transcript])


def _vector_path(video_id: str) -> str:
    return str(VECTOR_DIR / video_id)


def create_vector_store(docs, video_id: str):
    embedding = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    persist_dir = _vector_path(video_id)
    if Path(persist_dir).exists() and any(Path(persist_dir).iterdir()):
        return Chroma(persist_directory=persist_dir, embedding_function=embedding)
    return Chroma.from_documents(docs, embedding, persist_directory=persist_dir)


def rag_answer(question: str, vectorstore) -> str:
    results = vectorstore.similarity_search(question, k=4)
    context_text = "\n".join(i.page_content for i in results)
    return _run(
        """You are a kind, polite, precise assistant.
        Answer ONLY using the context. If the answer is not present, say so and ask the
        user to rephrase.

        Context:
        {context}

        Question:
        {question}

        Answer:""",
        context=context_text,
        question=question,
    )


# ---------- Exporters ----------

def generate_mindmap_structure(transcript: str) -> dict:
    """Ask the LLM for a hierarchical mind-map JSON: central -> branches -> sub-points."""
    raw = _run(
        """Extract a mind-map of this transcript. Output STRICT JSON only — no prose, no
        code fences. Schema:
        {{
          "central": "Short central topic (2-5 words)",
          "branches": [
            {{
              "title": "Major theme (2-5 words)",
              "children": ["Sub-point 1", "Sub-point 2", "Sub-point 3"]
            }}
          ]
        }}
        Use 4-7 branches, each with 2-5 children. Children should be concise (under 10 words).

        Transcript:
        {transcript}""",
        transcript=transcript,
    )
    cleaned = raw.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?\s*|\s*```$", "", cleaned, flags=re.MULTILINE).strip()
    try:
        return json.loads(cleaned)
    except Exception:
        m = re.search(r"\{.*\}", cleaned, re.DOTALL)
        if m:
            return json.loads(m.group(0))
        raise


def render_mindmap_html(structure: dict) -> str:
    """Build a pyvis interactive mind map and return its HTML as a string."""
    from pyvis.network import Network

    net = Network(height="650px", width="100%", bgcolor="#0e1117", font_color="white", directed=False)
    net.barnes_hut(gravity=-8000, spring_length=180)

    central = structure.get("central", "Video")
    net.add_node(
        "root", label=central, color="#ff6b6b", size=40,
        shape="box", font={"size": 22, "color": "white", "face": "arial"},
    )

    palette = ["#4ecdc4", "#ffe66d", "#a8e6cf", "#ff8b94", "#c7ceea", "#ffd3b6", "#b5ead7"]
    for i, branch in enumerate(structure.get("branches", [])):
        bid = f"b{i}"
        color = palette[i % len(palette)]
        net.add_node(
            bid, label=branch.get("title", f"Topic {i+1}"),
            color=color, size=28, shape="box",
            font={"size": 16, "color": "#0e1117"},
        )
        net.add_edge("root", bid, color=color, width=3)
        for j, child in enumerate(branch.get("children", [])):
            cid = f"b{i}c{j}"
            net.add_node(
                cid, label=child, color="#ffffff", size=18, shape="dot",
                font={"size": 13, "color": "white"},
            )
            net.add_edge(bid, cid, color=color, width=1.5)

    return net.generate_html(notebook=False)


def export_markdown(title: str, content: str) -> bytes:
    return f"# {title}\n\n{content}\n".encode("utf-8")


def export_docx(title: str, content: str) -> bytes:
    from docx import Document
    from io import BytesIO

    doc = Document()
    doc.add_heading(title, level=1)
    for line in content.split("\n"):
        stripped = line.strip()
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith(("- ", "* ")):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped:
            doc.add_paragraph(stripped)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_pdf(title: str, content: str) -> bytes:
    from fpdf import FPDF
    from io import BytesIO

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Helvetica", "B", 16)
    safe_title = title.encode("latin-1", "replace").decode("latin-1")
    pdf.multi_cell(0, 10, safe_title)
    pdf.ln(2)
    pdf.set_font("Helvetica", size=11)
    for line in content.split("\n"):
        safe_line = line.encode("latin-1", "replace").decode("latin-1")
        pdf.multi_cell(0, 6, safe_line)
    return bytes(pdf.output(dest="S"))
